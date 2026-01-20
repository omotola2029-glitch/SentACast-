import streamlit as st
import pandas as pd
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
from prophet import Prophet
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import re
from io import BytesIO
from collections import Counter
import time

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Inches as DocxInches, Pt as DocxPt, RGBColor as DocxRGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from wordcloud import WordCloud
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import numpy as np

st.set_page_config(
    page_title="SentACast - Sentiment Analysis & Forecasting",
    page_icon="sentacast.png",
    layout="wide"
)

# Custom CSS 
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }
    
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        font-weight: 600;
    }
    
    .stButton>button {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        background-color: #28a745 !important;
        color: white !important;
        border: none !important;
        padding: 0.5rem 1rem !important;
        height: auto !important;
        min-height: 38px !important;
        max-height: 45px !important;
        line-height: 1.2 !important;
    }
    
    .stButton>button:hover {
        background-color: #218838 !important;
        color: white !important;
    }
    
    .stDownloadButton>button {
        background-color: #007bff !important;
        color: white !important;
        padding: 0.5rem 1rem !important;
        height: auto !important;
        min-height: 38px !important;
        max-height: 45px !important;
    }
    
    .stDownloadButton>button:hover {
        background-color: #0056b3 !important;
    }
    
    .confidence-high {
        color: #28a745;
        font-weight: bold;
    }
    
    .confidence-medium {
        color: #ffc107;
        font-weight: bold;
    }
    
    .confidence-low {
        color: #dc3545;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

def generate_sample_data():
    """Generate 50 sample clothing reviews for demo purposes."""
    sample_reviews = [
        ("2024-01-15", "This shirt fits perfectly and the fabric quality is amazing! Very comfortable to wear all day.", 5),
        ("2024-01-16", "The jeans are okay but the sizing runs a bit small. Quality could be better for the price.", 3),
        ("2024-01-17", "Absolutely love this dress! The color is vibrant and it looks exactly like the photos.", 5),
        ("2024-01-18", "Terrible quality. The sweater started pilling after just one wash. Very disappointed.", 1),
        ("2024-01-19", "Great value for money. The t-shirt is comfortable and the material feels durable.", 4),
        ("2024-01-20", "The jacket is stylish but the zipper broke within a week. Poor construction.", 2),
        ("2024-01-21", "Perfect fit! These pants are exactly what I was looking for. Highly recommend!", 5),
        ("2024-01-22", "Average quality. Nothing special but not terrible either. It's what you'd expect for the price.", 3),
        ("2024-01-23", "The fabric is so soft and comfortable! Best hoodie I've ever owned.", 5),
        ("2024-01-24", "Disappointed with the color. It looks completely different from the website photos.", 2),
        ("2024-01-25", "Excellent quality and fast shipping. The blouse fits great and looks professional.", 5),
        ("2024-01-26", "Too tight around the shoulders. Had to return it. Size chart was not accurate.", 2),
        ("2024-01-27", "Love the style and material! Very trendy and comfortable for casual wear.", 4),
        ("2024-01-28", "The stitching came apart after two wears. Very poor quality control.", 1),
        ("2024-02-02", "Perfect for the gym! The leggings don't slip and the material breathes well.", 5),
        ("2024-01-30", "It's okay. Nothing to write home about but serves its purpose.", 3),
        ("2024-01-31", "Fantastic quality! The coat is warm and looks premium. Worth every penny.", 5),
        ("2024-02-01", "Fabric feels cheap and scratchy. Not comfortable at all.", 2),
        ("2024-02-02", "Perfect for the gym! The leggings don't slip and the material breathes well.", 5),
        ("2024-02-03", "Runs very large. Had to exchange for a smaller size. Otherwise nice quality.", 3),
        ("2024-02-04", "The shirt is nice but wrinkles easily. Needs ironing after every wash.", 3),
        ("2024-02-05", "Extremely disappointed. The sweater shrank dramatically in the wash.", 1),
        ("2024-02-06", "Great fit and the color is even better in person! Very satisfied.", 4),
        ("2024-02-07", "Not worth the money. You can find better quality elsewhere for the same price.", 2),
        ("2024-02-08", "", 5),
        ("2024-02-09", "Decent purchase. The shorts are comfortable for lounging at home.", 3),
        ("2024-02-10", "The dress arrived damaged. Customer service was helpful with the refund though.", 2),
        ("2024-02-11", "Best purchase I've made this year! The quality exceeds expectations.", 5),
        ("2024-02-12", "Average fit. Nothing stands out as particularly good or bad.", 3),
        ("2024-02-13", "The material is so luxurious! Feels expensive and looks amazing.", 5),
        ("2024-02-14", "Horrible experience. Wrong item sent and it took weeks to get a replacement.", 1),
        ("2024-02-15", "Love it! Perfect for office wear. Professional and comfortable.", 4),
        ("2024-02-16", "The pants are too long even though I ordered my usual size.", 3),
        ("2024-02-17", "Excellent craftsmanship! You can tell this was made with care.", 5),
        ("2024-02-18", "Not as described. The fabric is much thinner than I expected.", 2),
        ("2024-02-19", "Great everyday wear! Comfortable and easy to style with anything.", 4),
        ("2024-02-20", "The seams are poorly done. Already starting to fray.", 2),
        ("2024-02-21", "Absolutely perfect! Fits like it was made for me. Love the style!", 5),
        ("2024-02-22", "It's fine. Does the job but nothing exciting about it.", 3),
        ("2024-02-23", "Beautiful quality and the details are lovely! Very impressed.", 5),
        ("2024-02-24", "Arrived late and the packaging was damaged. Product seems okay though.", 3),
        ("2024-02-25", "Fantastic! The sweater is cozy and looks great. Will buy more colors.", 5),
        ("2024-02-26", "Too expensive for what you get. Quality doesn't justify the price.", 2),
        ("2024-02-27", "Perfect summer top! Light, breathable, and stylish.", 4),
        ("2024-02-28", "The fit is weird. Too tight in some places and too loose in others.", 2),
        ("2024-02-29", "Incredible value! Can't believe the quality for this price point.", 5),
        ("2024-03-01", "Standard quality. Nothing to complain about but nothing to rave about either.", 3),
        ("2024-03-02", "Love the design! Very trendy and gets lots of compliments.", 4),
        ("2024-03-03", "Color faded after first wash. Very disappointing quality.", 2),
        ("2024-03-04", "Best clothing purchase I've made in a long time! Absolutely love it!", 5)
    ]
    
    df = pd.DataFrame(sample_reviews, columns=['date', 'review_text', 'rating'])
    df['date'] = pd.to_datetime(df['date'])
    return df

def clean_text(text):
    """Clean review text by removing URLs, special characters, and normalizing format."""
    if pd.isna(text) or text == "":
        return ""
    
    text = str(text)
    text = text.lower()
    text = re.sub(r'http\S+|www\S+|https\S+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\S+@\S+', '', text)
    text = re.sub(r'<.*?>', '', text)
    text = re.sub(r'@\w+', '', text)
    text = re.sub(r'#(\w+)', r'\1', text)
    text = re.sub(r'([!?.]){4,}', r'\1\1\1', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text

def calculate_confidence(pos, neg, neu, compound):
    """Calculate confidence score for VADER prediction."""
    # Confidence is higher when one polarity dominates
    max_polarity = max(pos, neg, neu)
    polarity_spread = max_polarity - min(pos, neg, neu)
    
    # Also consider absolute compound score
    compound_strength = abs(compound)
    
    # Combined confidence (0-1 scale)
    confidence = (polarity_spread * 0.6) + (compound_strength * 0.4)
    return min(confidence, 1.0)

def get_confidence_label(confidence):
    """Get confidence label and color."""
    if confidence >= 0.7:
        return "High", "confidence-high", "#28a745"
    elif confidence >= 0.4:
        return "Medium", "confidence-medium", "#ffc107"
    else:
        return "Low", "confidence-low", "#dc3545"

def generate_wordcloud(text_data, title, colormap='Blues'):
    """Generate word cloud from text data."""
    if len(text_data) == 0 or all(len(str(t).strip()) == 0 for t in text_data):
        return None
    
    # Combine all text
    combined_text = ' '.join(str(t) for t in text_data if pd.notna(t) and str(t).strip())
    
    if len(combined_text.strip()) == 0:
        return None
    # Generate word cloud
    wordcloud = WordCloud(
        width=800,
        height=400,
        background_color='#e8f4f8',
        colormap=colormap,
        max_words=50,
        relative_scaling=0.5,
        min_font_size=10
    ).generate(combined_text)
    # Create matplotlib figure
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.set_title(title, fontsize=16, fontweight='bold')
    ax.axis('off')
    plt.tight_layout(pad=0)
    
    return fig

def save_plot_to_image(fig):
    """Save matplotlib figure to bytes for embedding in reports."""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

def save_plotly_to_image(fig):
    """Save plotly figure to bytes for embedding in reports."""
    buf = BytesIO()
    fig.write_image(buf, format='png', width=800, height=600)
    buf.seek(0)
    return buf

def extract_emotions(df):
    """Extract emotion words from reviews for emotion analysis."""
    emotion_words = {
        'Satisfaction': ['satisfied', 'happy', 'pleased', 'content', 'glad', 'delighted'],
        'Disappointment': ['disappointed', 'frustrated', 'unhappy', 'dissatisfied', 'letdown'],
        'Excitement': ['excited', 'thrilled', 'amazing', 'awesome', 'love', 'fantastic'],
        'Anger': ['angry', 'furious', 'terrible', 'horrible', 'worst', 'hate'],
        'Gratitude': ['thanks', 'grateful', 'appreciate', 'thank you', 'thankful'],
        'Surprise': ['surprised', 'unexpected', 'wow', 'shocked', 'amazed'] }
    emotion_counts = {emotion: 0 for emotion in emotion_words.keys()}
    
    for text in df['review_text']:
        text_lower = str(text).lower()
        for emotion, words in emotion_words.items():
            if any(word in text_lower for word in words):
                emotion_counts[emotion] += 1
    
    return emotion_counts

def extract_intent(df):
    """Extract user intent from reviews."""
    intent_keywords = {
        'Praise': ['great', 'excellent', 'perfect', 'amazing', 'love', 'best', 'good', 'wonderful'],
        'Complaint': ['bad', 'poor', 'terrible', 'worst', 'disappointed', 'horrible', 'awful'],
        'Recommend': ['recommend', 'suggest', 'worth', 'should buy', 'must have'],
        'Question': ['?', 'how', 'why', 'what', 'when', 'where', 'wondering'],
        'Suggest': ['should', 'could', 'would be better', 'improve', 'needs', 'wish'],
        'Inform': ['arrived', 'received', 'came', 'got', 'delivered', 'shipped']}
    
    intent_counts = {intent: 0 for intent in intent_keywords.keys()}
    
    for text in df['review_text']:
        text_lower = str(text).lower()
        for intent, keywords in intent_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                intent_counts[intent] += 1
    
    return intent_counts

def preprocess_dataframe(df):
    """Preprocess the dataframe with automatic column detection and data cleaning."""
    preprocessing_stats = {
        'original_rows': len(df),
        'removed_duplicates': 0,
        'removed_empty': 0,
        'removed_invalid_dates': 0,
        'detected_columns': {}
    }
    
    # detect review text column
    possible_review_columns = [
        'review_text', 'review text', 'reviewtext',
        'review_body', 'review body', 'reviewbody',
        'text', 'body', 'content', 'comment', 'feedback',
        'review', 'customer_review'
    ]
    
    review_col = None
    for col in df.columns:
        if col.lower().strip() in possible_review_columns:
            review_col = col
            preprocessing_stats['detected_columns']['review_text'] = col
            break
    
    if review_col is None:
        available_cols = ', '.join(df.columns.tolist())
        raise ValueError(f"No review text column found. Available columns: {available_cols}")
    
    if review_col != 'review_text':
        df = df.rename(columns={review_col: 'review_text'})
    
    # detect date column
    possible_date_columns = [
        'date', 'review_date', 'review date', 'timestamp',
        'time', 'datetime', 'created_date', 'posted_date'
    ]
    
    date_col = None
    for col in df.columns:
        if col.lower().strip() in possible_date_columns:
            date_col = col
            preprocessing_stats['detected_columns']['date'] = col
            break
    
    if date_col and date_col != 'date':
        df = df.rename(columns={date_col: 'date'})
    
    # detect rating column
    possible_rating_columns = [
        'rating', 'star_rating', 'star rating', 'stars',
        'star', 'score', 'review_rating'
    ]
    
    rating_col = None
    for col in df.columns:
        if col.lower().strip() in possible_rating_columns:
            rating_col = col
            preprocessing_stats['detected_columns']['rating'] = col
            break
    
    if rating_col and rating_col != 'rating':
        df = df.rename(columns={rating_col: 'rating'})
    
    if 'rating' in df.columns:
        df['rating'] = pd.to_numeric(df['rating'], errors='coerce')
    
    # Remove duplicates
    original_len = len(df)
    df = df.drop_duplicates(subset=['review_text'], keep='first')
    preprocessing_stats['removed_duplicates'] = original_len - len(df)
    
    # Save original text
    df['original_text'] = df['review_text'].copy()
    
    # Clean text
    df['review_text'] = df['review_text'].apply(clean_text)
    
    # Remove empty reviews
    original_len = len(df)
    df = df[df['review_text'].str.strip() != '']
    preprocessing_stats['removed_empty'] = original_len - len(df)
    
    # Processdates
    if 'date' in df.columns:
        original_len = len(df)
        df['date'] = pd.to_datetime(df['date'], errors='coerce')
        df = df.dropna(subset=['date'])
        preprocessing_stats['removed_invalid_dates'] = original_len - len(df)
        df = df.sort_values('date')
    # Calculate text statistics
    df['text_length'] = df['review_text'].str.len()
    df['word_count'] = df['review_text'].str.split().str.len()
    df = df.reset_index(drop=True)
    
    preprocessing_stats['final_rows'] = len(df)
    
    return df, preprocessing_stats

def calculate_comprehensive_metrics(df, stage="analysis"):
    """Calculate all performance metrics from the analyzed dataframe."""
    metrics = {}
    
    metrics['total_reviews'] = len(df)
    
    if 'sentiment' in df.columns:
        sentiment_counts = df['sentiment'].value_counts()
        metrics['positive_count'] = sentiment_counts.get('Positive', 0)
        metrics['negative_count'] = sentiment_counts.get('Negative', 0)
        metrics['neutral_count'] = sentiment_counts.get('Neutral', 0)
        
        metrics['positive_pct'] = (metrics['positive_count'] / metrics['total_reviews']) * 100
        metrics['negative_pct'] = (metrics['negative_count'] / metrics['total_reviews']) * 100
        metrics['neutral_pct'] = (metrics['neutral_count'] / metrics['total_reviews']) * 100
        
        metrics['avg_compound'] = df['compound'].mean()
        metrics['median_compound'] = df['compound'].median()
        metrics['std_compound'] = df['compound'].std()
        metrics['min_compound'] = df['compound'].min()
        metrics['max_compound'] = df['compound'].max()
        
        metrics['avg_positive_score'] = df['pos'].mean()
        metrics['avg_negative_score'] = df['neg'].mean()
        metrics['avg_neutral_score'] = df['neu'].mean()
        
        if 'confidence' in df.columns:
            metrics['avg_confidence'] = df['confidence'].mean()
            metrics['high_confidence_count'] = (df['confidence'] >= 0.7).sum()
            metrics['low_confidence_count'] = (df['confidence'] < 0.4).sum()
        
        metrics['emotions'] = extract_emotions(df)
        metrics['intent'] = extract_intent(df)
    
    if 'text_length' in df.columns:
        metrics['avg_text_length'] = df['text_length'].mean()
        metrics['median_text_length'] = df['text_length'].median()
        metrics['min_text_length'] = df['text_length'].min()
        metrics['max_text_length'] = df['text_length'].max()
    
    if 'word_count' in df.columns:
        metrics['avg_word_count'] = df['word_count'].mean()
        metrics['median_word_count'] = df['word_count'].median()
    
    if 'needs_review' in df.columns:
        metrics['flagged_count'] = df['needs_review'].sum()
        metrics['flagging_rate'] = (metrics['flagged_count'] / metrics['total_reviews']) * 100
    
    if 'corrected' in df.columns:
        metrics['corrected_count'] = df['corrected'].sum()
        metrics['correction_rate'] = (metrics['corrected_count'] / metrics['total_reviews']) * 100
    
    if 'date' in df.columns:
        metrics['date_range_start'] = df['date'].min()
        metrics['date_range_end'] = df['date'].max()
        metrics['date_span_days'] = (metrics['date_range_end'] - metrics['date_range_start']).days
        metrics['reviews_per_day'] = metrics['total_reviews'] / max(metrics['date_span_days'], 1)
    
    return metrics

def display_metrics_dashboard(metrics, title="System Metrics"):
    """Display comprehensive metrics dashboard with visualizations."""
    st.subheader(title)
    
    st.divider()
    st.markdown("### Overview")
    col1, col2, col3, col4 = st.columns(4)
    
    col1.metric("Total Reviews", f"{metrics['total_reviews']:,}")
    
    if 'avg_compound' in metrics:
        col2.metric("Average Sentiment", f"{metrics['avg_compound']:.3f}")
        col3.metric("Sentiment Range", f"{metrics['min_compound']:.2f} to {metrics['max_compound']:.2f}")
        if 'avg_confidence' in metrics:
            col4.metric("Avg Confidence", f"{metrics['avg_confidence']:.1%}")
        else:
            col4.metric("Std Deviation", f"{metrics['std_compound']:.3f}")
    
    st.divider()
    
    if 'positive_count' in metrics:
        st.markdown("### Sentiment Distribution")
        col1, col2, col3 = st.columns(3)
        
        col1.metric("Positive Reviews", f"{metrics['positive_count']:,}", delta=f"{metrics['positive_pct']:.1f}%")
        col2.metric("Neutral Reviews", f"{metrics['neutral_count']:,}", delta=f"{metrics['neutral_pct']:.1f}%")
        col3.metric("Negative Reviews", f"{metrics['negative_count']:,}", delta=f"{metrics['negative_pct']:.1f}%")
    
    st.divider()
    
    if 'avg_positive_score' in metrics:
        st.markdown("### Average Polarity Scores")
        col1, col2, col3 = st.columns(3)
        
        col1.metric("Positive Score", f"{metrics['avg_positive_score']:.3f}")
        col2.metric("Negative Score", f"{metrics['avg_negative_score']:.3f}")
        col3.metric("Neutral Score", f"{metrics['avg_neutral_score']:.3f}")
    
    st.divider()
    
    # Emotions analysis
    if 'emotions' in metrics and sum(metrics['emotions'].values()) > 0:
        st.markdown("### Emotions Driving Product Quality")
        emotion_df = pd.DataFrame(list(metrics['emotions'].items()), columns=['Emotion', 'Count'])
        emotion_df = emotion_df[emotion_df['Count'] > 0].sort_values('Count', ascending=False)
        
        if len(emotion_df) > 0:
            fig_emotions = px.bar(
                emotion_df,
                x='Emotion',
                y='Count',
                title='Emotion Distribution in Reviews',
                color='Count',
                color_continuous_scale='Blues',
                labels={'Count': 'Number of Reviews'})
            fig_emotions.update_layout(showlegend=False)
            st.plotly_chart(fig_emotions, use_container_width=True)
    
    st.divider()
    
    # intention Analysis
    if 'intent' in metrics and sum(metrics['intent'].values()) > 0:
        st.markdown("### Intent of User Contributing to Product")
        intent_df = pd.DataFrame(list(metrics['intent'].items()), columns=['Intent', 'Count'])
        intent_df = intent_df[intent_df['Count'] > 0].sort_values('Count', ascending=False)
        
        if len(intent_df) > 0:
            fig_intent = px.pie(
                intent_df,
                names='Intent',
                values='Count',
                title='User Intent Distribution',
                color_discrete_sequence=px.colors.qualitative.Set3
            )
            st.plotly_chart(fig_intent, use_container_width=True)
    
    st.divider()
    
    if 'avg_text_length' in metrics:
        st.markdown("### Text Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        col1.metric("Avg Length (chars)", f"{metrics['avg_text_length']:.0f}")
        col2.metric("Median Length", f"{metrics['median_text_length']:.0f}")
        col3.metric("Avg Word Count", f"{metrics.get('avg_word_count', 0):.0f}")
        col4.metric("Length Range", f"{metrics['min_text_length']}-{metrics['max_text_length']}")
    
    st.divider()
    
    if 'flagged_count' in metrics:
        st.markdown("### Human-in-the-Loop Metrics")
        col1, col2, col3 = st.columns(3)
        
        col1.metric("Reviews Flagged", f"{metrics['flagged_count']:,}", delta=f"{metrics['flagging_rate']:.1f}%")
        
        if 'corrected_count' in metrics:
            col2.metric("Human Corrections", f"{metrics['corrected_count']:,}", delta=f"{metrics['correction_rate']:.1f}%")
            
            if metrics['flagged_count'] > 0:
                correction_effectiveness = (metrics['corrected_count'] / metrics['flagged_count']) * 100
                col3.metric("Correction Rate", f"{correction_effectiveness:.1f}%")
    
    st.divider()
    
    if 'date_range_start' in metrics:
        st.markdown("### Time Period Analysis")
        col1, col2, col3 = st.columns(3)
        
        col1.metric("Date Range", f"{metrics['date_range_start'].date()} to {metrics['date_range_end'].date()}")
        col2.metric("Total Days", f"{metrics['date_span_days']:,}")
        col3.metric("Reviews per Day", f"{metrics['reviews_per_day']:.1f}")

def create_summary_statistics_doc(df, metrics):
    """Generate text-based summary statistics document for export."""
    summary = []
    summary.append("=" * 80)
    summary.append("SENTACAST - SENTIMENT ANALYSIS SYSTEM SUMMARY")
    summary.append("=" * 80)
    summary.append(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    summary.append("\n" + "=" * 80)
    
    summary.append("\n1. OVERVIEW")
    summary.append("-" * 80)
    summary.append(f"Total Reviews Analyzed: {metrics['total_reviews']:,}")
    
    if 'avg_compound' in metrics:
        summary.append(f"Average Sentiment Score: {metrics['avg_compound']:.4f}")
        summary.append(f"Median Sentiment Score: {metrics['median_compound']:.4f}")
        summary.append(f"Standard Deviation: {metrics['std_compound']:.4f}")
        summary.append(f"Sentiment Range: {metrics['min_compound']:.4f} to {metrics['max_compound']:.4f}")
        
        if 'avg_confidence' in metrics:
            summary.append(f"Average Confidence: {metrics['avg_confidence']:.2%}")
    
    if 'positive_count' in metrics:
        summary.append("\n2. SENTIMENT DISTRIBUTION")
        summary.append("-" * 80)
        summary.append(f"Positive Reviews: {metrics['positive_count']:,} ({metrics['positive_pct']:.2f}%)")
        summary.append(f"Neutral Reviews: {metrics['neutral_count']:,} ({metrics['neutral_pct']:.2f}%)")
        summary.append(f"Negative Reviews: {metrics['negative_count']:,} ({metrics['negative_pct']:.2f}%)")
    
    if 'emotions' in metrics:
        summary.append("\n3. EMOTIONS ANALYSIS")
        summary.append("-" * 80)
        for emotion, count in sorted(metrics['emotions'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                summary.append(f"{emotion}: {count}")
    
    if 'intent' in metrics:
        summary.append("\n4. USER INTENT ANALYSIS")
        summary.append("-" * 80)
        for intent, count in sorted(metrics['intent'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                summary.append(f"{intent}: {count}")
    
    if 'avg_positive_score' in metrics:
        summary.append("\n5. AVERAGE POLARITY SCORES")
        summary.append("-" * 80)
        summary.append(f"Average Positive Score: {metrics['avg_positive_score']:.4f}")
        summary.append(f"Average Negative Score: {metrics['avg_negative_score']:.4f}")
        summary.append(f"Average Neutral Score: {metrics['avg_neutral_score']:.4f}")
    
    if 'avg_text_length' in metrics:
        summary.append("\n6. TEXT STATISTICS")
        summary.append("-" * 80)
        summary.append(f"Average Text Length: {metrics['avg_text_length']:.2f} characters")
        summary.append(f"Median Text Length: {metrics['median_text_length']:.2f} characters")
        summary.append(f"Average Word Count: {metrics.get('avg_word_count', 0):.2f} words")
        summary.append(f"Length Range: {metrics['min_text_length']} - {metrics['max_text_length']} characters")
    
    if 'flagged_count' in metrics:
        summary.append("\n7. HUMAN-IN-THE-LOOP METRICS")
        summary.append("-" * 80)
        summary.append(f"Reviews Flagged for Review: {metrics['flagged_count']:,} ({metrics['flagging_rate']:.2f}%)")
        
        if 'corrected_count' in metrics:
            summary.append(f"Human Corrections Applied: {metrics['corrected_count']:,} ({metrics['correction_rate']:.2f}%)")
            if metrics['flagged_count'] > 0:
                correction_effectiveness = (metrics['corrected_count'] / metrics['flagged_count']) * 100
                summary.append(f"Correction Effectiveness: {correction_effectiveness:.2f}%")
    
    if 'date_range_start' in metrics:
        summary.append("\n8. TIME PERIOD ANALYSIS")
        summary.append("-" * 80)
        summary.append(f"Date Range: {metrics['date_range_start'].date()} to {metrics['date_range_end'].date()}")
        summary.append(f"Total Days: {metrics['date_span_days']:,}")
        summary.append(f"Average Reviews per Day: {metrics['reviews_per_day']:.2f}")
    
    summary.append("\n" + "=" * 80)
    summary.append("END OF REPORT")
    summary.append("=" * 80)
    
    return "\n".join(summary)

def generate_pdf_report(df, metrics, forecast_data=None, sentiment_chart=None, intent_chart=None):
    """Generate comprehensive PDF report with embedded charts."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#28a745'),
        spaceAfter=30,
        alignment=TA_CENTER)
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#007bff'),
        spaceAfter=12,
        spaceBefore=12)
    
    title = Paragraph("SentACast Sentiment Analysis Report", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    date_text = Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
    story.append(date_text)
    story.append(Spacer(1, 20))
    
    # Summary
    story.append(Paragraph("1. Summary", heading_style))
    summary_text = f"""
    This report gives a comprehensive sentiment analysis of {metrics['total_reviews']:,} customer reviews.
    The average sentiment score is {metrics.get('avg_compound', 0):.3f}, indicating 
    {"positive" if metrics.get('avg_compound', 0) > 0.05 else "negative" if metrics.get('avg_compound', 0) < -0.05 else "neutral"} 
    overall customer sentiment.
    """
    story.append(Paragraph(summary_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Sentiment Distribution Table
    if 'positive_count' in metrics:
        story.append(Paragraph("2. Sentiment Distribution", heading_style))
        
        data = [
            ['Sentiment', 'Count', 'Percentage'],
            ['Positive', f"{metrics['positive_count']:,}", f"{metrics['positive_pct']:.2f}%"],
            ['Neutral', f"{metrics['neutral_count']:,}", f"{metrics['neutral_pct']:.2f}%"],
            ['Negative', f"{metrics['negative_count']:,}", f"{metrics['negative_pct']:.2f}%"],
            ['Total', f"{metrics['total_reviews']:,}", "100.00%"]
        ]
        
        table = Table(data, colWidths=[2*inch, 2*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#28a745')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
        # sentiment chart
        if sentiment_chart:
            story.append(Paragraph("Sentiment Distribution Visualization", heading_style))
            story.append(RLImage(sentiment_chart, width=5*inch, height=3.5*inch))
            story.append(Spacer(1, 20))
    
    # Emotions Analysis
    if 'emotions' in metrics:
        story.append(Paragraph("3. Emotions Analysis", heading_style))
        emotion_text = "Emotions detected in customer reviews:<br/><br/>"
        for emotion, count in sorted(metrics['emotions'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                emotion_text += f"&bull; <b>{emotion}:</b> {count} reviews<br/>"
        story.append(Paragraph(emotion_text, styles['Normal']))
        story.append(Spacer(1, 12))
    # Intention Analysis
    if 'intent' in metrics:
        story.append(Paragraph("4. User Intent Analysis", heading_style))
        intent_text = "Customer intentions identified in reviews:<br/><br/>"
        for intent, count in sorted(metrics['intent'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                intent_text += f"&bull; <b>{intent}:</b> {count} reviews<br/>"
        story.append(Paragraph(intent_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # intent chart 
        if intent_chart:
            story.append(Paragraph("Intent Distribution Visualization", heading_style))
            story.append(RLImage(intent_chart, width=5*inch, height=3.5*inch))
            story.append(Spacer(1, 20))
    
    # metrics
    story.append(Paragraph("5. Key Performance Metrics", heading_style))
    if 'avg_compound' in metrics:
        metrics_text = f"""
        <b>Sentiment Metrics:</b><br/>
        Average Sentiment Score: {metrics['avg_compound']:.4f}<br/>
        Median Sentiment Score: {metrics['median_compound']:.4f}<br/>
        Standard Deviation: {metrics['std_compound']:.4f}<br/>
        """
        if 'avg_confidence' in metrics:
            metrics_text += f"Average Confidence: {metrics['avg_confidence']:.2%}<br/>"
        
        metrics_text += f"""
        <br/>
        <b>Text Statistics:</b><br/>
        Average Review Length: {metrics.get('avg_text_length', 0):.0f} characters<br/>
        Average Word Count: {metrics.get('avg_word_count', 0):.0f} words<br/>
        """
        story.append(Paragraph(metrics_text, styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_word_report(df, metrics, forecast_data=None, sentiment_chart=None, intent_chart=None):
    """Generate comprehensive Word document report with embedded charts."""
    doc = Document()
    
    # Title
    title = doc.add_heading('SentACast Sentiment Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('1. Summary', 1)
    doc.add_paragraph(
        f"In this report is a comprehensive sentiment analysis of {metrics['total_reviews']:,} customer reviews. "
        f"The average sentiment score is {metrics.get('avg_compound', 0):.3f}, indicating "
        f"{'positive' if metrics.get('avg_compound', 0) > 0.05 else 'negative' if metrics.get('avg_compound', 0) < -0.05 else 'neutral'} "
        f"overall customer sentiment."
    )
    doc.add_paragraph()
    
    # sentiment distribution
    if 'positive_count' in metrics:
        doc.add_heading('2. Sentiment Distribution', 1)
        
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sentiment'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        
        table.rows[1].cells[0].text = 'Positive'
        table.rows[1].cells[1].text = f"{metrics['positive_count']:,}"
        table.rows[1].cells[2].text = f"{metrics['positive_pct']:.2f}%"
        
        table.rows[2].cells[0].text = 'Neutral'
        table.rows[2].cells[1].text = f"{metrics['neutral_count']:,}"
        table.rows[2].cells[2].text = f"{metrics['neutral_pct']:.2f}%"
        
        table.rows[3].cells[0].text = 'Negative'
        table.rows[3].cells[1].text = f"{metrics['negative_count']:,}"
        table.rows[3].cells[2].text = f"{metrics['negative_pct']:.2f}%"
        
        table.rows[4].cells[0].text = 'Total'
        table.rows[4].cells[1].text = f"{metrics['total_reviews']:,}"
        table.rows[4].cells[2].text = "100.00%"
        
        doc.add_paragraph()
        
        # sentiment chart
        if sentiment_chart:
            doc.add_paragraph('Sentiment Distribution Visualization:', style='Heading 3')
            doc.add_picture(sentiment_chart, width=DocxInches(5.5))
            doc.add_paragraph()
    
    # Emotions Analysis
    if 'emotions' in metrics:
        doc.add_heading('3. Emotions Analysis', 1)
        doc.add_paragraph("Emotions detected in customer reviews:")
        
        for emotion, count in sorted(metrics['emotions'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                doc.add_paragraph(f"{emotion}: {count} reviews", style='List Bullet')
        
        doc.add_paragraph()
    
    # Intention analysis
    if 'intent' in metrics:
        doc.add_heading('4. User Intent Analysis', 1)
        doc.add_paragraph("Customer intentions identified in reviews:")
        
        for intent, count in sorted(metrics['intent'].items(), key=lambda x: x[1], reverse=True):
            if count > 0:
                doc.add_paragraph(f"{intent}: {count} reviews", style='List Bullet')
        
        doc.add_paragraph()       
        #intent chart 
        if intent_chart:
            doc.add_paragraph('Intent Distribution Visualization:', style='Heading 3')
            doc.add_picture(intent_chart, width=DocxInches(5.5))
            doc.add_paragraph()
    
    # metrics
    doc.add_heading('5. Key Performance Metrics', 1)
    if 'avg_compound' in metrics:
        doc.add_heading('Sentiment Metrics:', 2)
        doc.add_paragraph(f"Average Sentiment Score: {metrics['avg_compound']:.4f}")
        doc.add_paragraph(f"Median Sentiment Score: {metrics['median_compound']:.4f}")
        doc.add_paragraph(f"Standard Deviation: {metrics['std_compound']:.4f}")
        
        if 'avg_confidence' in metrics:
            doc.add_paragraph(f"Average Confidence: {metrics['avg_confidence']:.2%}")
        
        doc.add_heading('Text Statistics:', 2)
        doc.add_paragraph(f"Average Review Length: {metrics.get('avg_text_length', 0):.0f} characters")
        doc.add_paragraph(f"Average Word Count: {metrics.get('avg_word_count', 0):.0f} words")
    
    # hitl metrics
    if 'flagged_count' in metrics:
        doc.add_heading('6. Human-in-the-Loop Metrics', 1)
        doc.add_paragraph(f"Reviews Flagged for Review: {metrics['flagged_count']:,}")
        if 'corrected_count' in metrics:
            doc.add_paragraph(f"Human Corrections Applied: {metrics['corrected_count']:,}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# initialize session state
if 'raw_df' not in st.session_state:
    st.session_state['raw_df'] = None
if 'cleaned_df' not in st.session_state:
    st.session_state['cleaned_df'] = None
if 'analyzed_df' not in st.session_state:
    st.session_state['analyzed_df'] = None
if 'corrected_df' not in st.session_state:
    st.session_state['corrected_df'] = None
if 'preprocessing_stats' not in st.session_state:
    st.session_state['preprocessing_stats'] = None
if 'forecast_data' not in st.session_state:
    st.session_state['forecast_data'] = None
if 'start_time' not in st.session_state:
    st.session_state['start_time'] = time.time()
if 'vader_threshold' not in st.session_state:
    st.session_state['vader_threshold'] = 0.05
if 'forecast_days' not in st.session_state:
    st.session_state['forecast_days'] = 30
if 'hitl_enabled' not in st.session_state:
    st.session_state['hitl_enabled'] = True
if 'data_cleaned' not in st.session_state:
    st.session_state['data_cleaned'] = False
if 'data_analyzed' not in st.session_state:
    st.session_state['data_analyzed'] = False

# the title and logo
col1, col2 = st.columns([1, 5])
with col1:
        st.image("sentacast.png", width=700)

with col2:
    st.title("SentACast")
    st.markdown("*Sentiment Analysis and Forecasting System with Human Review*")

st.divider()

# Sidebar nav
st.sidebar.header("Navigation")
stage = st.sidebar.radio(
    "Select Stage:",
    ["Home", "Upload and Clean Data", "Analyze Sentiment", "Review and Correction", "Forecast Trends", "Metrics Dashboard"])

st.sidebar.divider()

# ettings  (collapsible
with st.sidebar.expander(" Settings", expanded=False):
    st.markdown("**System Configuration:**")
    
    vader_threshold = st.slider(
        "VADER Threshold",
        min_value=0.01,
        max_value=0.20,
        value=st.session_state['vader_threshold'],
        step=0.01,
        help="Sensitivity for sentiment classification"
    )
    st.session_state['vader_threshold'] = vader_threshold
    
    forecast_days_setting = st.slider(
        "Default Forecast Days",
        min_value=7,
        max_value=90,
        value=st.session_state['forecast_days'],
        step=7
    )
    st.session_state['forecast_days'] = forecast_days_setting
    
    hitl_toggle = st.checkbox(
        "Enable HITL",
        value=st.session_state['hitl_enabled'],
        help="Enable Human-in-the-Loop review"
    )
    st.session_state['hitl_enabled'] = hitl_toggle

# session Info
st.sidebar.divider()
st.sidebar.markdown("🟢 **Status:** All Systems Ready")

# calc session time
session_minutes = int((time.time() - st.session_state['start_time']) / 60)
st.sidebar.markdown(f"⏱️ **Session:** {session_minutes} min")

# dataset nfo 
if st.session_state.get('cleaned_df') is not None:
    st.sidebar.divider()
    st.sidebar.markdown("**Current Dataset:**")
    df_info = st.session_state['cleaned_df']
    st.sidebar.markdown(f"• Reviews: {len(df_info):,}")
    if 'date' in df_info.columns:
        st.sidebar.markdown(f"• Date Range: {df_info['date'].min().date()}")

st.sidebar.divider()
if st.sidebar.button("🔄 Reset All Data", use_container_width=True):
    for key in ['raw_df', 'cleaned_df', 'analyzed_df', 'corrected_df', 'preprocessing_stats', 'forecast_data', 'data_cleaned', 'data_analyzed']:
        st.session_state[key] = None if key in ['raw_df', 'cleaned_df', 'analyzed_df', 'corrected_df', 'preprocessing_stats', 'forecast_data'] else False
    st.session_state['start_time'] = time.time()
    st.success("All data reset successfully!")
    st.rerun()

# main

if stage == "Home":
    st.header("Welcome to SentACast")
    
    st.markdown("---")
    
    st.markdown("### About SentACast")
    st.write("""
    SentACast is an AI-driven sentiment analysis and forecasting system designed specifically for e-commerce platforms. 
    The system combines automated sentiment analysis with human oversight to ensure accurate and reliable results for business decision-making.
    """)
    
    st.markdown("---")
    
    st.markdown("### Key Features")
    st.write("""
    SentACast provides comprehensive sentiment analysis abilities including automatic data cleaning and preprocessing, 
    VADER sentiment analysis for accurate classification of product reviews, emotion and intent extraction to understand customer 
    motivations, human-in-the-loop correction mechanisms for more accuracy, sentiment trend forecasting using Prophet for 
    prediction, comprehensive metrics and visualizations for detailed analysis, and export in many formats  
    including PDF reports, Word documents, CSV data files, and text summaries.
    """)
    
    st.markdown("---")
    
    st.markdown("### Getting Started")
    st.write("""
    To begin your analysis, upload a CSV file containing customer reviews. The file should include a column with the review text 
    such as review_text, review body, or something similar. If time-based forecasting is needed, include a date column
    and optionally a rating column for additional insights. SentACast will automatically detects and process these columns.
    """)
    
    with st.expander("View Example Data Format"):
        st.code("""
date,review_text,rating
2024-01-15,This product exceeded my expectations,5
2024-01-16,Quality could be better for the price,3
2024-01-17,Absolutely love it! Highly recommend,5
        """, language="csv")
    
    st.markdown("---")
    st.info("👈 Use the Navigation menu on the left to begin your analysis")

elif stage == "Upload and Clean Data":
    st.header("Upload Reviews and Clean Data")
    
    
    
    st.info("""
    Upload your CSV file containing product reviews. The system will automatically detect review text, date, and rating columns, 
    remove duplicates and empty entries, clean text by removing URLs and special characters, and provide detailed statistics on the cleaning process.
    """)
    
    # data sample
    col1, col2 = st.columns([1, 1])
    with col1:
        uploaded_file = st.file_uploader(
            "Upload your product reviews CSV file",
            type=['csv'],
            help="File must contain a review text column"
        )
    
    with col2:
        st.write("")
        st.write("")
        if st.button("Load Sample Data (Clothing Reviews)", type="secondary", use_container_width=True):
            sample_df = generate_sample_data()
            st.session_state['raw_df'] = sample_df
            st.success("Sample clothing reviews loaded! Click 'Clean and Preprocess Data' below.")
            st.rerun()
    
    if uploaded_file or st.session_state.get('raw_df') is not None:
        try:
            if uploaded_file and st.session_state['raw_df'] is None:
                raw_df = pd.read_csv(uploaded_file)
                st.session_state['raw_df'] = raw_df
            else:
                raw_df = st.session_state['raw_df']
            
            st.success(f"File loaded successfully. Found {len(raw_df)} reviews")
            
            st.markdown("---")
            
            with st.expander("View Raw Data (First 5 Reviews)"):
                st.dataframe(raw_df.head())
            
            st.markdown("---")
            
            if st.button("Clean & Preprocess Data", type="primary", key="clean_data") or st.session_state.get('data_cleaned'):
                if not st.session_state.get('data_cleaned'):
                    with st.spinner("Cleaning and preprocessing data..."):
                        try:
                            cleaned_df, stats = preprocess_dataframe(raw_df.copy())
                            st.session_state['cleaned_df'] = cleaned_df
                            st.session_state['preprocessing_stats'] = stats
                            st.session_state['data_cleaned'] = True
                            st.success("Data cleaning complete")
                        except ValueError as e:
                            st.error(f"Error: {str(e)}")
                            st.stop()
                
                cleaned_df = st.session_state['cleaned_df']
                stats = st.session_state['preprocessing_stats']
                
                st.markdown("---")
                st.subheader("Preprocessing Metrics")
                
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("Original Reviews", f"{stats['original_rows']:,}")
                col2.metric("Duplicates Removed", f"{stats['removed_duplicates']:,}")
                col3.metric("Empty Removed", f"{stats['removed_empty']:,}")
                col4.metric("Final Clean Reviews", f"{stats['final_rows']:,}")
                
                retention_rate = (stats['final_rows'] / stats['original_rows']) * 100
                st.info(f"Data Retention Rate: {retention_rate:.1f}% of original data retained")
                
                st.markdown("---")
                
                if 'date' in cleaned_df.columns:
                    col1, col2 = st.columns(2)
                    start_date = str(cleaned_df['date'].min().date())
                    end_date = str(cleaned_df['date'].max().date())
                    col1.metric("Start Date", start_date)
                    col2.metric("End Date", end_date)
                    st.markdown("---")
                
                st.subheader("Text Statistics After Cleaning")
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("Avg Length", f"{cleaned_df['text_length'].mean():.0f} chars")
                col2.metric("Median Length", f"{cleaned_df['text_length'].median():.0f} chars")
                col3.metric("Min Length", f"{cleaned_df['text_length'].min()} chars")
                col4.metric("Max Length", f"{cleaned_df['text_length'].max()} chars")
                
                st.markdown("---")
                
                fig_length = px.histogram(
                    cleaned_df,
                    x='text_length',
                    nbins=30,
                    title='Review Length Distribution After Cleaning',
                    labels={'text_length': 'Character Count'},
                    color_discrete_sequence=['#17becf']
                )
                st.plotly_chart(fig_length, use_container_width=True)
                
                st.markdown("---")
                
                csv = cleaned_df.to_csv(index=False)
                st.download_button(
                    label="Download Cleaned CSV",
                    data=csv,
                    file_name=f"cleaned_reviews_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    key="download_cleaned")
                st.info("👈 Next: Select 'Analyze Sentiment' from the navigation menu")
        
        except Exception as e:
            st.error(f"Error loading file: {str(e)}")

elif stage == "Analyze Sentiment":
    st.header("Sentiment Analysis")
    
    if st.session_state['cleaned_df'] is None:
        st.warning("Please complete Upload and Clean Data first")
        st.stop()
    
    df = st.session_state['cleaned_df'].copy()
    
    st.success(f"Using cleaned dataset with {len(df):,} reviews")
       
    if st.button("Run Sentiment Analysis", type="primary", key="run_vader") or st.session_state.get('data_analyzed'):
        if not st.session_state.get('data_analyzed'):
            with st.spinner("Analyzing sentiments..."):
                analyzer = SentimentIntensityAnalyzer()
                
                def get_vader_scores(text):
                    if pd.isna(text) or text == "":
                        return pd.Series({'compound': 0, 'pos': 0, 'neg': 0, 'neu': 1})
                    scores = analyzer.polarity_scores(str(text))
                    return pd.Series({
                        'compound': scores['compound'],
                        'pos': scores['pos'],
                        'neg': scores['neg'],
                        'neu': scores['neu']
                    })
                
                df[['compound', 'pos', 'neg', 'neu']] = df['review_text'].apply(get_vader_scores)
                
     # calc confidence scores
                df['confidence'] = df.apply(lambda row: calculate_confidence(row['pos'], row['neg'], row['neu'], row['compound']), axis=1)
                
                threshold = st.session_state['vader_threshold']
                def classify_sentiment(compound):
                    if compound >= threshold:
                        return "Positive"
                    elif compound <= -threshold:
                        return "Negative"
                    else:
                        return "Neutral"
                
                df['sentiment'] = df['compound'].apply(classify_sentiment)
                
                # flag if HITL enabled
                if st.session_state['hitl_enabled']:
                    def flag_for_review(row):
                        reasons = []
                        text = str(row['review_text']).lower()
                        
                        if abs(row['compound']) < threshold:
                            reasons.append("Ambiguous sentiment")
                        
                        if row['pos'] > 0.3 and row['neg'] > 0.3:
                            reasons.append("Mixed emotions detected")
                        
                        # low confidence flagging
                        if row['confidence'] < 0.4:
                            reasons.append("Low confidence prediction")
                        
                        sarcasm_words = ['yeah right', 'sure', 'obviously', 'wicked']
                        if any(word in text for word in sarcasm_words) and '!' in text:
                            reasons.append("Possible sarcasm")
                        
                        if 'not' in text or "n't" in text or 'no' in text:
                            positive_words = ['good', 'great', 'excellent', 'amazing', 'love']
                            if any(word in text for word in positive_words):
                                reasons.append("Negation detected")
                        
                        return '; '.join(reasons) if reasons else None
                    
                    df['flag_reason'] = df.apply(flag_for_review, axis=1)
                    df['needs_review'] = df['flag_reason'].notna()
                else:
                    df['flag_reason'] = None
                    df['needs_review'] = False
                
                df['corrected'] = False
                
                st.session_state['analyzed_df'] = df
                st.session_state['data_analyzed'] = True
                st.success("Analysis Complete")
        
        df = st.session_state['analyzed_df']
        
        st.markdown("---")
        
        metrics = calculate_comprehensive_metrics(df)
        display_metrics_dashboard(metrics, "Sentiment Analysis Metrics")
        
        st.markdown("---")
        
        st.subheader("Sentiment Visualizations")
        
        sentiment_counts = df['sentiment'].value_counts()
        fig_pie = px.pie(
            values=sentiment_counts.values,
            names=sentiment_counts.index,
            title="Sentiment Distribution",
            color=sentiment_counts.index,
            color_discrete_map={
                'Positive': '#28a745',
                'Neutral': '#ffc107',
                'Negative': '#dc3545'
            }
        )
        st.plotly_chart(fig_pie, use_container_width=True)
        
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        
        with col1:
            fig_hist = px.histogram(
                df,
                x='compound',
                nbins=30,
                title='Compound Score Distribution',
                labels={'compound': 'Compound Score'},
                color_discrete_sequence=['#1f77b4']
            )
            st.plotly_chart(fig_hist, use_container_width=True)
        
        st.markdown("---")
        
        # wordcloud
        st.subheader("Word Cloud")
        st.write("Most frequently used words in each sentiment category:")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("**Positive Reviews**")
            positive_text = df[df['sentiment'] == 'Positive']['review_text']
            if len(positive_text) > 0:
                wc_fig = generate_wordcloud(positive_text, '', 'Greens')
                if wc_fig:
                    st.pyplot(wc_fig)
            else:
                st.info("No positive reviews found")
        
        with col2:
            st.markdown("**Neutral Reviews**")
            neutral_text = df[df['sentiment'] == 'Neutral']['review_text']
            if len(neutral_text) > 0:
                wc_fig = generate_wordcloud(neutral_text, '', 'Blues')
                if wc_fig:
                    st.pyplot(wc_fig)
            else:
                st.info("No neutral reviews found")
        
        with col3:
            st.markdown("**Negative Reviews**")
            negative_text = df[df['sentiment'] == 'Negative']['review_text']
            if len(negative_text) > 0:
                wc_fig = generate_wordcloud(negative_text, '', 'Reds')
                if wc_fig:
                    st.pyplot(wc_fig)
            else:
                st.info("No negative reviews found")
        
        st.markdown("---")
        
        # prediction confidence
        st.subheader("Prediction Confidence")
        col1, col2, col3 = st.columns(3)
        
        high_conf = (df['confidence'] >= 0.7).sum()
        med_conf = ((df['confidence'] >= 0.4) & (df['confidence'] < 0.7)).sum()
        low_conf = (df['confidence'] < 0.4).sum()
        
        col1.metric("High Confidence", f"{high_conf:,}", delta=f"{(high_conf/len(df)*100):.1f}%")
        col2.metric("Medium Confidence", f"{med_conf:,}", delta=f"{(med_conf/len(df)*100):.1f}%")
        col3.metric("Low Confidence", f"{low_conf:,}", delta=f"{(low_conf/len(df)*100):.1f}%")
        
        fig_confidence = px.histogram(
            df,
            x='confidence',
            nbins=20,
            title='Confidence Score Distribution',
            labels={'confidence': 'Confidence Score'},
            color_discrete_sequence=['#17becf']
        )
        st.plotly_chart(fig_confidence, use_container_width=True)
        
        st.markdown("---")
        
        st.subheader("Flagged Reviews")
        flag_data = pd.DataFrame({
            'Status': ['Needs Review', 'Approved'],
            'Count': [df['needs_review'].sum(), (~df['needs_review']).sum()]
        })
        
        fig_flag = px.bar(
            flag_data,
            x='Status',
            y='Count',
            title='Review Status',
            color='Status',
            color_discrete_map={'Needs Review': '#dc3545', 'Approved': '#28a745'}
        )
        st.plotly_chart(fig_flag, use_container_width=True)
        
        st.markdown("---")
        
        fig_scatter = px.scatter(
            df,
            x='text_length',
            y='compound',
            color='sentiment',
            title='Review Length vs Sentiment',
            labels={'text_length': 'Review Length (characters)', 'compound': 'Sentiment Score'},
            color_discrete_map={
                'Positive': '#28a745',
                'Neutral': '#ffc107',
                'Negative': '#dc3545'
            }
        )
        st.plotly_chart(fig_scatter, use_container_width=True)
        
        st.markdown("---")
        
        # download csv analysed reviews
        st.subheader("Download Analysis Results")
        csv_analyzed = df.to_csv(index=False)
        st.download_button(
            label="Download Analyzed Reviews CSV",
            data=csv_analyzed,
            file_name=f"analyzed_reviews_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=False,
            key="download_analyzed"
        )
        st.info("👈 Next: Select 'Review and Correction' from the navigation menu")

elif stage == "Review and Correction":
    st.header("Human Review and Correction")
    
    if st.session_state['analyzed_df'] is None:
        st.warning("Please complete Analyze Sentiment first")
        st.stop()
    
    df = st.session_state['analyzed_df'].copy()
    flagged_df = df[df['needs_review']].copy()
    
    if len(flagged_df) == 0:
        st.success("No reviews flagged for correction")
        st.info("All sentiment classifications look reliable.")
        st.session_state['corrected_df'] = df
        
        st.markdown("---")
        st.info("👈 Next: Select 'Forecast Trends' from the navigation menu")
    else:
        st.write(f"**{len(flagged_df):,} reviews flagged for human review**")
        
        st.markdown("---")
        
        st.subheader("HITL Metrics Before Corrections")
        col1, col2, col3 = st.columns(3)
        col1.metric("Total Reviews", f"{len(df):,}")
        col2.metric("Flagged for Review", f"{len(flagged_df):,}")
        col3.metric("Flagging Rate", f"{(len(flagged_df)/len(df)*100):.1f}%")
        
        st.info("Review each flagged item and decide if VADER classified it correctly. If not, select the correct sentiment.")
        
        corrections = {}
        
        for i, (idx, row) in enumerate(flagged_df.iterrows(), 1):
            with st.expander(f"Review {i} of {len(flagged_df)} - {row['flag_reason']}", expanded=(i==1)):
                
                col1, col2, col3 = st.columns([3, 1.5, 1.5])
                
                with col1:
                    st.markdown("**Cleaned Review:**")
                    st.text_area(
                        "Text",
                        row['review_text'],
                        height=100,
                        disabled=True,
                        key=f"text_{idx}",
                        label_visibility="collapsed"
                    )
                    
                    if 'original_text' in row and row['original_text'] != row['review_text']:
                        with st.expander("See original before cleaning"):
                            st.caption(row['original_text'])
                    
                    st.caption(f"Flag Reason: {row['flag_reason']}")
                    st.caption(f"Length: {row['text_length']} characters | Words: {row['word_count']}")
                
                with col2:
                    st.markdown("**VADER Scores:**")
                    st.metric("Compound", f"{row['compound']:.3f}")
                    st.write(f"Positive: {row['pos']:.2f}")
                    st.write(f"Negative: {row['neg']:.2f}")
                    st.write(f"Neutral: {row['neu']:.2f}")
                    
                    # Show confidence
                    conf_label, conf_class, conf_color = get_confidence_label(row['confidence'])
                    st.markdown(f"**Confidence:** <span class='{conf_class}'>{conf_label} ({row['confidence']:.1%})</span>", unsafe_allow_html=True)
                    
                    if row['compound'] >= 0.05:
                        st.success(f"{row['sentiment']}")
                    elif row['compound'] <= -0.05:
                        st.error(f"{row['sentiment']}")
                    else:
                        st.warning(f"{row['sentiment']}")
                
                with col3:
                    st.markdown("**Your Decision:**")
                    correction = st.selectbox(
                        "Correct sentiment to:",
                        ["Keep VADER Result", "Positive", "Neutral", "Negative"],
                        key=f"correct_{idx}",
                        label_visibility="collapsed"
                    )
                    
                    if correction != "Keep VADER Result":
                        correction_map = {
                            "Positive": 0.6,
                            "Neutral": 0.0,
                            "Negative": -0.6
                        }
                        corrections[idx] = correction_map[correction]
                        st.success("Will be corrected")
        
        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            if st.button("Apply Corrections and Continue", type="primary", use_container_width=True, key="apply_corrections"):
                if len(corrections) > 0:
                    for idx, new_score in corrections.items():
                        df.at[idx, 'compound'] = new_score
                        df.at[idx, 'corrected'] = True
                        
                        if new_score >= 0.05:
                            df.at[idx, 'sentiment'] = "Positive"
                        elif new_score <= -0.05:
                            df.at[idx, 'sentiment'] = "Negative"
                        else:
                            df.at[idx, 'sentiment'] = "Neutral"
                    
                    st.session_state['corrected_df'] = df
                    st.success(f"Applied {len(corrections):,} corrections successfully")
                    
                    correction_rate = (len(corrections) / len(flagged_df)) * 100
                    
                
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Reviews Corrected", f"{len(corrections):,}")
                    col2.metric("Correction Rate", f"{correction_rate:.1f}%")
                    col3.metric("Accuracy Improvement", f"+{correction_rate:.1f}%")
                else:
                    st.session_state['corrected_df'] = df
                    st.info("No corrections made. Proceeding with original VADER results.")
                
                st.markdown("---")
                st.info("👈 Next: Select 'Forecast Trends' from the navigation menu")

elif stage == "Forecast Trends":
    st.header("Sentiment Trend Forecasting")
    
    if st.session_state['corrected_df'] is not None:
        df = st.session_state['corrected_df']
    elif st.session_state['analyzed_df'] is not None:
        df = st.session_state['analyzed_df']
    else:
        st.warning("Please complete previous steps first")
        st.stop()
    

    
    if 'date' not in df.columns:
        st.error("Your CSV file needs a date column for forecasting")
        st.info("Format: YYYY-MM-DD (e.g., 2024-01-15)")
        st.info("You can still view sentiment analysis results in the Metrics Dashboard and export reports.")
        
        st.markdown("---")
        st.info("👈 View complete analysis: Select 'Metrics Dashboard' from the navigation menu")
        st.stop()
    
    ts_data = df.groupby('date')['compound'].agg(['mean', 'count']).reset_index()
    ts_data.columns = ['ds', 'y', 'review_count']
    
    st.subheader("Historical Sentiment Trend")

    col1, col2, col3, col4 = st.columns(4)
    start_date_str = str(ts_data['ds'].min().date())
    end_date_str = str(ts_data['ds'].max().date())
    
    col1.markdown(f"<p style='font-size:14px'><b>Date Range:</b><br>{start_date_str} to {end_date_str}</p>", unsafe_allow_html=True)
    col2.metric("Total Days", f"{len(ts_data):,}")
    col3.metric("Avg Daily Sentiment", f"{ts_data['y'].mean():.3f}")
    col4.metric("Total Reviews", f"{df.shape[0]:,}")
    
    fig_hist = go.Figure()
    fig_hist.add_trace(go.Scatter(
        x=ts_data['ds'],
        y=ts_data['y'],
        mode='lines+markers',
        name='Average Sentiment',
        line=dict(color='#1f77b4', width=2),
        marker=dict(size=6)
    ))
    fig_hist.update_layout(
        title='Historical Sentiment Over Time',
        xaxis_title='Date',
        yaxis_title='Average Compound Score',
        hovermode='x unified'
    )
    st.plotly_chart(fig_hist, use_container_width=True)
    
   
    
    st.subheader("Generate Forecast")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        forecast_days = st.slider(
            "Select forecast period (days):",
            min_value=7,
            max_value=90,
            value=st.session_state['forecast_days'],
            step=7
        )
    
    with col2:
        st.write("")
        st.write("")
        forecast_button = st.button("Generate Forecast", type="primary", use_container_width=True, key="gen_forecast")
    
    if forecast_button or st.session_state.get('forecast_data') is not None:
        if forecast_button:
            with st.spinner("Running Prophet forecasting model..."):
                model = Prophet(
                    daily_seasonality=False,
                    weekly_seasonality=True,
                    yearly_seasonality=False,
                    interval_width=0.95
                )
                model.fit(ts_data[['ds', 'y']])
                
                future = model.make_future_dataframe(periods=forecast_days)
                forecast = model.predict(future)
                
                future_forecast = forecast[forecast['ds'] > ts_data['ds'].max()]
                st.session_state['forecast_data'] = future_forecast
                st.session_state['full_forecast'] = forecast
                st.session_state['forecast_model'] = model
                
                st.success("Forecast generated successfully")
        
        future_forecast = st.session_state['forecast_data']
        forecast = st.session_state.get('full_forecast')
        model = st.session_state.get('forecast_model')
        
        st.markdown("---")
        
        st.subheader("Forecast Results")
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric(
            "Predicted Avg Sentiment",
            f"{future_forecast['yhat'].mean():.3f}",
            delta=f"{future_forecast['yhat'].mean() - ts_data['y'].mean():.3f}"
        )
        col2.metric(
            "Forecast Trend", 
            "Improving" if future_forecast['yhat'].iloc[-1] > future_forecast['yhat'].iloc[0] 
            else "Declining"
        )
        col3.metric("Confidence Interval", "95%")
        col4.metric("Forecast Period", f"{len(future_forecast)} days")
        
        
        
        fig_forecast = go.Figure()
        
        fig_forecast.add_trace(go.Scatter(
            x=ts_data['ds'],
            y=ts_data['y'],
            mode='markers',
            name='Historical Data',
            marker=dict(color='#1f77b4', size=8)
        ))
        
        fig_forecast.add_trace(go.Scatter(
            x=forecast['ds'],
            y=forecast['yhat'],
            mode='lines',
            name='Forecast',
            line=dict(color='#28a745', width=3)
        ))
        
        fig_forecast.add_trace(go.Scatter(
            x=forecast['ds'],
            y=forecast['yhat_upper'],
            mode='lines',
            line=dict(width=0),
            showlegend=False,
            hoverinfo='skip'
        ))
        
        fig_forecast.add_trace(go.Scatter(
            x=forecast['ds'],
            y=forecast['yhat_lower'],
            mode='lines',
            line=dict(width=0),
            fill='tonexty',
            fillcolor='rgba(40, 167, 69, 0.2)',
            name='Confidence Interval',
            hoverinfo='skip'
        ))
        
        fig_forecast.update_layout(
            title=f'Sentiment Forecast - Next {len(future_forecast)} Days',
            xaxis_title='Date',
            yaxis_title='Predicted Sentiment Score',
            hovermode='x unified',
            height=500
        )
        
        st.plotly_chart(fig_forecast, use_container_width=True)
        
        st.subheader("Forecast Components")
        st.write("Understanding what drives the sentiment trends")
        
        if model:
            fig_components = model.plot_components(forecast)
            st.pyplot(fig_components)
        
        st.markdown("---")
        
        st.subheader("Download Forecast Data")
        
        forecast_export = future_forecast[['ds', 'yhat', 'yhat_lower', 'yhat_upper']].copy()
        forecast_export.columns = ['Date', 'Predicted_Sentiment', 'Lower_Bound', 'Upper_Bound']
        
        csv_forecast = forecast_export.to_csv(index=False)
        st.download_button(
            label="Download Forecast CSV",
            data=csv_forecast,
            file_name=f"sentiment_forecast_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=False,
            key="download_forecast_csv"
        )
        st.info("👈 View complete analysis: Select 'Metrics Dashboard' from the navigation menu")

elif stage == "Metrics Dashboard":
    st.header("Comprehensive Metrics Dashboard")
    
    if st.session_state['corrected_df'] is not None:
        df = st.session_state['corrected_df']
        st.success("Displaying metrics for corrected data (after HITL)")
    elif st.session_state['analyzed_df'] is not None:
        df = st.session_state['analyzed_df']
        st.info("Displaying metrics for analyzed data (before HITL corrections)")
    elif st.session_state['cleaned_df'] is not None:
        df = st.session_state['cleaned_df']
        st.warning("Only preprocessing metrics available. Complete sentiment analysis first.")
    else:
        st.warning("No data available. Please complete the pipeline first.")
        st.stop()
    
    # Dashboard summary
    if 'sentiment' in df.columns:
        st.subheader("Dashboard Summary")
        
        metrics = calculate_comprehensive_metrics(df)
        
        # metrics in highlighted boxes
        col1, col2, col3, col4 = st.columns(4)
        
        sentiment_status = "Positive" if metrics['avg_compound'] > 0.05 else "Negative" if metrics['avg_compound'] < -0.05 else "Neutral"
        sentiment_color = "🟢" if sentiment_status == "Positive" else "🔴" if sentiment_status == "Negative" else "🟡"
        
        col1.metric(
            "Overall Sentiment",
            f"{sentiment_color} {sentiment_status}",
            delta=f"{metrics['avg_compound']:.3f}"
        )
        col2.metric(
            "Customer Satisfaction",
            f"{metrics['positive_pct']:.1f}%",
            delta=f"{metrics['positive_count']:,} reviews"
        )
        col3.metric(
            "Total Reviews Analyzed",
            f"{metrics['total_reviews']:,}",
            delta=f"{metrics.get('avg_confidence', 0):.0%} avg confidence" if 'avg_confidence' in metrics else ""
        )
        
        if 'date_range_start' in metrics:
            col4.metric(
                "Analysis Period",
                f"{metrics['date_span_days']} days",
                delta=f"{metrics['reviews_per_day']:.1f} per day"
            )
        
        st.divider()
    
    # Full metrics dashboard
    metrics = calculate_comprehensive_metrics(df)
    display_metrics_dashboard(metrics, "Complete System Performance Metrics")
    
    if 'corrected' in df.columns and df['corrected'].sum() > 0:
        st.divider()
        st.subheader("Before vs After HITL Corrections")
        
        original_df = st.session_state.get('analyzed_df')
        if original_df is not None:
            st.markdown("---")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### Before HITL")
                before_sentiment = original_df['sentiment'].value_counts()
                before_avg = original_df['compound'].mean()
                
                st.metric("Average Sentiment", f"{before_avg:.3f}")
                st.write(f"Positive: {before_sentiment.get('Positive', 0):,}")
                st.write(f"Neutral: {before_sentiment.get('Neutral', 0):,}")
                st.write(f"Negative: {before_sentiment.get('Negative', 0):,}")
            
            with col2:
                st.markdown("### After HITL")
                after_sentiment = df['sentiment'].value_counts()
                after_avg = df['compound'].mean()
                
                st.metric("Average Sentiment", f"{after_avg:.3f}", 
                         delta=f"{after_avg - before_avg:.3f}")
                st.write(f"Positive: {after_sentiment.get('Positive', 0):,}")
                st.write(f"Neutral: {after_sentiment.get('Neutral', 0):,}")
                st.write(f"Negative: {after_sentiment.get('Negative', 0):,}")
            
            st.markdown("---")
            
            comparison_data = pd.DataFrame({
                'Stage': ['Before HITL'] * 3 + ['After HITL'] * 3,
                'Sentiment': ['Positive', 'Neutral', 'Negative'] * 2,
                'Count': [
                    before_sentiment.get('Positive', 0),
                    before_sentiment.get('Neutral', 0),
                    before_sentiment.get('Negative', 0),
                    after_sentiment.get('Positive', 0),
                    after_sentiment.get('Neutral', 0),
                    after_sentiment.get('Negative', 0)
                ]
            })
            
            fig_comparison = px.bar(
                comparison_data,
                x='Stage',
                y='Count',
                color='Sentiment',
                barmode='group',
                title='Sentiment Distribution: Before vs After HITL',
                color_discrete_map={
                    'Positive': '#28a745',
                    'Neutral': '#ffc107',
                    'Negative': '#dc3545'
                }
            )
            st.plotly_chart(fig_comparison, use_container_width=True)
    
    st.divider()
    st.subheader("Export Complete Reports")
    
    st.markdown("---")
    
    # chart images for reports
    sentiment_chart_img = None
    intent_chart_img = None
    
    if 'sentiment' in df.columns:
        # sentiment pie chart
        sentiment_counts = df['sentiment'].value_counts()
        fig_sentiment = px.pie(
            values=sentiment_counts.values,
            names=sentiment_counts.index,
            title="Sentiment Distribution",
            color=sentiment_counts.index,
            color_discrete_map={
                'Positive': '#28a745',
                'Neutral': '#ffc107',
                'Negative': '#dc3545'
            }
        )
        try:
            sentiment_chart_img = save_plotly_to_image(fig_sentiment)
        except:
            pass  
        
        # intent pie chart
        intent_counts = extract_intent(df)
        intent_df = pd.DataFrame(list(intent_counts.items()), columns=['Intent', 'Count'])
        intent_df = intent_df[intent_df['Count'] > 0]
        
        if len(intent_df) > 0:
            fig_intent = px.pie(
                intent_df,
                names='Intent',
                values='Count',
                title='User Intent Distribution'
            )
            try:
                intent_chart_img = save_plotly_to_image(fig_intent)
            except:
                pass
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        pdf_buffer = generate_pdf_report(df, metrics, st.session_state.get('forecast_data'), sentiment_chart_img, intent_chart_img)
        st.download_button(
            label="Download PDF Report",
            data=pdf_buffer,
            file_name=f"sentacast_report_{datetime.now().strftime('%Y%m%d')}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="metrics_pdf"
        )
    
    with col2:
        docx_buffer = generate_word_report(df, metrics, st.session_state.get('forecast_data'), sentiment_chart_img, intent_chart_img)
        st.download_button(
            label="Download Word Report",
            data=docx_buffer,
            file_name=f"sentacast_report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="metrics_docx"
        )
    
    with col3:
        summary_text = create_summary_statistics_doc(df, metrics)
        st.download_button(
            label="Download Summary Statistics",
            data=summary_text,
            file_name=f"sentacast_summary_{datetime.now().strftime('%Y%m%d')}.txt",
            mime="text/plain",
            use_container_width=True,
            key="metrics_summary"
        )
